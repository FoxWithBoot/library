from docx import Document


def create_doc(tab):
    document = Document()
    document.add_heading('Отчет по книгам', 0)
    table = document.add_table(rows=1, cols=4)
    table.rows[0].cells[0].text = 'Издание'
    table.rows[0].cells[1].text = 'Кол-во'
    table.rows[0].cells[2].text = 'Книга'
    table.rows[0].cells[3].text = 'Востребованность'
    for i in tab:
        cells = table.add_row().cells
        cells[0].text = i.title
        cells[1].text = str(i.count if i.count else 0)
        cells[2].text = i.name if i.name else 'НЕТ'
        cells[3].text = str(i.v_c if i.v_c else 0)
    table.style = 'LightShading-Accent1'
    return document

